"""
Automated Report Generator

Generates Word/HTML reports from modeling results.
"""
import base64
import io
import os
from datetime import datetime
from typing import Any, Dict, List, Optional

import pandas as pd


def generate_html_report(
    modeling_result: Dict[str, Any],
    data_info: Optional[Dict] = None,
    shap_result: Optional[Dict] = None,
    quality_report: Optional[Dict] = None,
) -> str:
    """Generate an HTML report."""
    lines = [
        '<!DOCTYPE html><html><head><meta charset="utf-8"><title>Modeling Report</title>',
        '<style>body{font-family:Arial,sans-serif;max-width:900px;margin:40px auto;line-height:1.6;color:#333}',
        'h1,h2{color:#2c3e50;border-bottom:2px solid #3498db;padding-bottom:6px}',
        'table{border-collapse:collapse;width:100%;margin:12px 0}',
        'th,td{border:1px solid #ddd;padding:8px;text-align:left;font-size:13px}',
        'th{background:#f8f9fa}tr:nth-child(even){background:#fafafa}',
        '.metric{display:inline-block;padding:8px 16px;background:#e8f4f8;border-radius:6px;margin:4px;font-size:14px}',
        'pre{background:#f4f4f4;padding:12px;border-radius:6px;overflow:auto;font-size:12px}',
        '</style></head><body>',
        f'<h1>Modeling Report</h1><p>Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>',
    ]
    
    # Data Summary
    if data_info:
        lines.append('<h2>1. Data Summary</h2>')
        lines.append(f'<p>Shape: {data_info.get("shape", "N/A")}</p>')
        if 'columns' in data_info:
            lines.append('<table><thead><tr><th>Column</th><th>Type</th></tr></thead><tbody>')
            for c in data_info['columns']:
                lines.append(f'<tr><td>{c}</td><td>-</td></tr>')
            lines.append('</tbody></table>')
    
    # Model Results
    if modeling_result:
        lines.append('<h2>2. Model Results</h2>')
        lines.append('<div>')
        for k, v in modeling_result.items():
            if isinstance(v, (int, float, str)):
                lines.append(f'<span class="metric">{k}: {v}</span>')
        lines.append('</div>')
        
        if 'leaderboard' in modeling_result and modeling_result['leaderboard']:
            lines.append('<h3>Leaderboard</h3><table><thead><tr>')
            df = pd.DataFrame(modeling_result['leaderboard'])
            for col in df.columns:
                lines.append(f'<th>{col}</th>')
            lines.append('</tr></thead><tbody>')
            for _, row in df.iterrows():
                lines.append('<tr>' + ''.join(f'<td>{v}</td>' for v in row) + '</tr>')
            lines.append('</tbody></table>')
    
    # SHAP
    if shap_result and 'feature_importance' in shap_result:
        lines.append('<h2>3. SHAP Feature Importance</h2><table><thead><tr><th>Feature</th><th>Importance</th></tr></thead><tbody>')
        for f in shap_result['feature_importance'][:15]:
            lines.append(f'<tr><td>{f["feature"]}</td><td>{f["importance"]:.6f}</td></tr>')
        lines.append('</tbody></table>')
    
    # Quality
    if quality_report:
        lines.append('<h2>4. Data Quality</h2>')
        lines.append(f'<p>Missing cells: {quality_report.get("missing_values", {}).get("total_missing_cells", 0)}</p>')
        lines.append(f'<p>Duplicate rows: {quality_report.get("duplicates", {}).get("duplicate_rows", 0)}</p>')
    
    lines.append('</body></html>')
    return '\n'.join(lines)


def generate_time_series_report(
    predictions: pd.DataFrame,
    date_col: str = 'date',
    target_col: str = 'target',
    id_cols: Optional[List[str]] = None,
    output_path: Optional[str] = None
) -> str:
    """生成时间序列预测的文字报告"""
    lines = []
    lines.append("=" * 60)
    lines.append("时间序列预测报告")
    lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("=" * 60)
    
    lines.append(f"\n数据概览:")
    lines.append(f"  预测记录数: {len(predictions)}")
    lines.append(f"  日期范围: {predictions[date_col].min()} ~ {predictions[date_col].max()}")
    if id_cols:
        lines.append(f"  序列数: {predictions[id_cols].drop_duplicates().shape[0]}")
    
    pred_col = [c for c in predictions.columns if 'pred' in c.lower()]
    if pred_col:
        lines.append(f"\n预测统计:")
        lines.append(f"  均值: {predictions[pred_col[0]].mean():.4f}")
        lines.append(f"  标准差: {predictions[pred_col[0]].std():.4f}")
        lines.append(f"  最小值: {predictions[pred_col[0]].min():.4f}")
        lines.append(f"  最大值: {predictions[pred_col[0]].max():.4f}")
    
    report_text = "\n".join(lines)
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report_text)
    
    return report_text


def generate_word_report(
    modeling_result: Dict[str, Any],
    data_info: Optional[Dict] = None,
    output_path: str = 'report.docx'
) -> str:
    """Generate a Word report. Falls back to HTML if python-docx not installed."""
    try:
        from docx import Document
        from docx.shared import Inches
        doc = Document()
        doc.add_heading('Modeling Report', 0)
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        if data_info:
            doc.add_heading('Data Summary', level=1)
            doc.add_paragraph(f'Shape: {data_info.get("shape", "N/A")}')
        
        if modeling_result:
            doc.add_heading('Model Results', level=1)
            for k, v in modeling_result.items():
                if isinstance(v, (int, float, str)):
                    doc.add_paragraph(f'{k}: {v}', style='List Bullet')
            if 'leaderboard' in modeling_result and modeling_result['leaderboard']:
                doc.add_heading('Leaderboard', level=2)
                table = doc.add_table(rows=1, cols=len(modeling_result['leaderboard'][0]))
                table.style = 'Light Grid Accent 1'
                for i, col in enumerate(modeling_result['leaderboard'][0].keys()):
                    table.rows[0].cells[i].text = str(col)
                for row in modeling_result['leaderboard']:
                    cells = table.add_row().cells
                    for i, val in enumerate(row.values()):
                        cells[i].text = str(val)
        
        doc.save(output_path)
        return output_path
    except ImportError:
        html_path = output_path.replace('.docx', '.html')
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(generate_html_report(modeling_result, data_info))
        return html_path
